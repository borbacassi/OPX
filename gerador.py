from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# ── CORES ──────────────────────────────────────────────────────────────────────
VERMELHO = RGBColor(0xFF, 0x00, 0x00)
VERDE_GRIFO  = "92D050"   # highlight color name não existe em python-docx direto; usamos XML
AMARELO_GRIFO = "FFFF00"


def set_highlight(run, cor_hex):
    """Aplica cor de destaque (highlight) a um run via XML."""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    # python-docx aceita nomes: yellow, green, cyan, magenta, blue, red, darkBlue, etc.
    cor_map = {
        "FFFF00": "yellow",
        "92D050": "green",
    }
    highlight.set(qn('w:val'), cor_map.get(cor_hex, "yellow"))
    rPr.append(highlight)


def paragrafo_vermelho(doc, texto, bold=False, tamanho=None, alinhamento=None):
    """Adiciona parágrafo com todo o texto em vermelho."""
    p = doc.add_paragraph()
    if alinhamento:
        p.alignment = alinhamento
    run = p.add_run(texto)
    run.bold = bold
    run.font.color.rgb = VERMELHO
    if tamanho:
        run.font.size = Pt(tamanho)
    return p


def paragrafo_misto(doc, partes):
    """
    Adiciona parágrafo com partes de formatação diferente.
    partes = lista de dicts: { 'texto': str, 'cor': RGBColor|None, 'bold': bool, 'highlight': str|None }
    """
    p = doc.add_paragraph()
    for parte in partes:
        run = p.add_run(parte.get('texto', ''))
        run.bold = parte.get('bold', False)
        cor = parte.get('cor')
        if cor:
            run.font.color.rgb = cor
        hl = parte.get('highlight')
        if hl:
            set_highlight(run, hl)
    return p


def formatar_data(valor):
    """Converte yyyy-mm-dd para dd/mm/aaaa se necessário."""
    if valor and len(valor) == 10 and valor[4] == '-':
        partes = valor.split('-')
        return f"{partes[2]}/{partes[1]}/{partes[0]}"
    return valor


def adicionar_bloco_rotulo(doc, rotulo, index):
    """Adiciona um bloco de rótulo ao documento."""

    nome_mat = rotulo.get('nome_material', '')
    tipo_imp = rotulo.get('tipo_impressao', '')
    qtd      = rotulo.get('qtd_rotulo', '')
    qtd_art  = rotulo.get('qtd_artes', '')
    val_tt   = rotulo.get('val_tt', '')
    val_und  = rotulo.get('val_und', '')
    obs      = rotulo.get('obs_arte', '').strip()
    artes    = rotulo.get('artes', [])   # lista de { 'imagem': path, 'larg': str, 'alt': str }

    # Linha: Nome Material + Tipo Impressão — grifado verde
    p_mat = doc.add_paragraph()
    run_mat = p_mat.add_run(f"{nome_mat}  {tipo_imp}")
    run_mat.bold = True
    set_highlight(run_mat, "92D050")

    # Linha: Qtd e artes — normal
    doc.add_paragraph(f"{qtd} unidades- ({qtd_art} {'ARTE' if qtd_art == '1' else 'ARTES'})")

    # Linha: Valores — normal
    doc.add_paragraph(f"Valores: R${val_tt}  - (R${val_und} por und)")

    # Observações — grifado amarelo
    if obs:
        p_obs = doc.add_paragraph()
        run_obs = p_obs.add_run(obs)
        set_highlight(run_obs, "FFFF00")

    # Artes com tamanhos
    num_artes = len(artes)

    if num_artes == 0:
        pass  # nenhuma arte enviada

    elif num_artes == 1:
        # Tamanho abaixo do nome do material, imagem abaixo
        arte = artes[0]
        larg = arte.get('larg', '')
        alt  = arte.get('alt', '')
        if larg or alt:
            doc.add_paragraph(f"{larg} x {alt}")
        img_path = arte.get('imagem', '')
        if img_path and os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.add_run().add_picture(img_path, width=Inches(2.5))

    else:
        # Múltiplas artes: cada imagem ao lado do seu tamanho
        # Usamos uma tabela invisível para alinhar lado a lado
        num_cols = num_artes
        tabela = doc.add_table(rows=1, cols=num_cols)
        tabela.style = 'Table Grid'

        # Remove bordas da tabela
        for row in tabela.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for lado in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{lado}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        for i, arte in enumerate(artes):
            cell = tabela.cell(0, i)
            larg = arte.get('larg', '')
            alt  = arte.get('alt', '')

            # Tamanho na célula
            if larg or alt:
                p_tam = cell.paragraphs[0]
                p_tam.add_run(f"{larg} x {alt}")
            
            # Imagem na célula
            img_path = arte.get('imagem', '')
            if img_path and os.path.exists(img_path):
                p_img = cell.add_paragraph()
                p_img.add_run().add_picture(img_path, width=Inches(1.8))

    doc.add_paragraph()  # espaço entre rótulos


def gerar_ordem(dados_gerais, rotulos, caminho_saida):
    doc = Document()

    # ── TÍTULO ────────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run("ORDEM DE PEDIDO")
    run_t.bold = True
    run_t.font.size = Pt(14)

    doc.add_paragraph()

    # ── CABEÇALHO: cliente / data / vendedor ──────────────────────────────────
    data_fmt = formatar_data(dados_gerais.get('data_atend', ''))
    p_cab = doc.add_paragraph()
    p_cab.add_run(
        f"{dados_gerais.get('nome_cliente', '')} – Pedido Whatsapp "
        f"{data_fmt} ({dados_gerais.get('nome_vendedor', '')})"
    )

    doc.add_paragraph()

    # ── TIPO DE ENVIO E PRAZOS — vermelho ─────────────────────────────────────
    d7  = formatar_data(dados_gerais.get('d7uteis', ''))
    d10 = formatar_data(dados_gerais.get('d10uteis', ''))
    d3  = formatar_data(dados_gerais.get('d3uteis', ''))
    d5  = formatar_data(dados_gerais.get('d5uteis', ''))

    p_env = doc.add_paragraph()
    p_env.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_env = p_env.add_run(
        f"{dados_gerais.get('tipo_envio', '')} DE {d7} A {d10}"
    )
    run_env.font.color.rgb = VERMELHO

    p_dig = doc.add_paragraph()
    run_dig = p_dig.add_run(f"AMOSTRA DIGITAL ATÉ {d3}")
    run_dig.font.color.rgb = VERMELHO

    p_fis = doc.add_paragraph()
    run_fis = p_fis.add_run(f"AMOSTRA FÍSICA ATÉ {d5}")
    run_fis.font.color.rgb = VERMELHO

    doc.add_paragraph()

    # ── RÓTULOS ───────────────────────────────────────────────────────────────
    for i, rotulo in enumerate(rotulos):
        adicionar_bloco_rotulo(doc, rotulo, i)

    # ── ENDEREÇOS ─────────────────────────────────────────────────────────────
    mesmo = dados_gerais.get('mesmo_endereco', False)

    if mesmo:
        doc.add_paragraph()
        p_tit = doc.add_paragraph()
        run_tit = p_tit.add_run("DADOS DO RECEBEDOR/NF:")
        run_tit.bold = True
        run_tit.font.color.rgb = VERMELHO

        campos = [
            ('Nome',           'nm_ENF'),
            ('CPF ou CNPJ',    'ID_ENF'),
            ('CEP',            'cep_ENF'),
            ('Rua ou Avenida', 'rua_ENF'),
            ('Número',         'num_end_ENF'),
            ('Complemento',    'comp_ENF'),
            ('Bairro',         'bairro_ENF'),
            ('Cidade/Estado',  'local_ENF'),
        ]
        for label, key in campos:
            doc.add_paragraph(f"{label}: {dados_gerais.get(key, '')}")

    else:
        # NF
        doc.add_paragraph()
        p_nf = doc.add_paragraph()
        run_nf = p_nf.add_run("DADOS DA NF:")
        run_nf.bold = True
        run_nf.font.color.rgb = VERMELHO

        for label, key in [
            ('Nome','nm_NF'),('CPF ou CNPJ','ID_NF'),('CEP','cep_NF'),
            ('Rua ou Avenida','rua_NF'),('Número','num_end_NF'),
            ('Complemento','comp_NF'),('Bairro','bairro_NF'),('Cidade/Estado','local_NF'),
        ]:
            doc.add_paragraph(f"{label}: {dados_gerais.get(key, '')}")

        # Entrega
        doc.add_paragraph()
        p_ent = doc.add_paragraph()
        run_ent = p_ent.add_run("DADOS DO RECEBEDOR:")
        run_ent.bold = True
        run_ent.font.color.rgb = VERMELHO

        for label, key in [
            ('Nome','nm_ENT'),('CPF ou CNPJ','ID_ENT'),('CEP','cep_ENT'),
            ('Rua ou Avenida','rua_ENT'),('Número','num_end_ENT'),
            ('Complemento','comp_ENT'),('Bairro','bairro_ENT'),('Cidade/Estado','local_ENT'),
        ]:
            doc.add_paragraph(f"{label}: {dados_gerais.get(key, '')}")

    doc.save(caminho_saida)
    print(f"Documento gerado: {caminho_saida}")